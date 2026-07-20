from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT = r"C:\Git\Safeviate-Manager-Vercel\formal_hearing_outcome_letter.docx"


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Finding Letter")
r.bold = True
r.font.name = "Arial"
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(31, 41, 55)
title.paragraph_format.space_after = Pt(10)

meta = doc.add_table(rows=2, cols=2)
meta.autofit = False
meta.columns[0].width = Inches(2.0)
meta.columns[1].width = Inches(4.9)
rows = [
    ("To:", "[Employee Name]"),
    ("Date:", "[Insert date]"),
]
for row, (label, value) in zip(meta.rows, rows):
    left, right = row.cells
    left.text = label
    right.text = value
    for cell in row.cells:
        set_cell_margins(cell)
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(11)
    left.paragraphs[0].runs[0].bold = True
    set_cell_shading(left, "EAF0F6")

doc.add_paragraph("")

salutation = doc.add_paragraph("Dear [Employee Name],")
salutation.paragraph_format.space_after = Pt(10)

body_paragraphs = [
    "This letter serves as the formal outcome of the disciplinary hearing held on [hearing date] concerning the allegation that you used drugs on 11 July 2026 and that a subsequent drug test conducted on 14 July 2026 returned a positive result.",
    "The chairperson considered the evidence presented during the hearing, including your admission that you used drugs on 11 July 2026, the urine test result obtained on 14 July 2026 indicating positive results for Cocaine and THC, and your decision not to proceed with blood or laboratory testing.",
    "Having assessed the evidence on the balance of probabilities, the chairperson finds that the allegation is proven.",
    "This conduct is a serious breach of workplace rules and is incompatible with the standards of conduct expected of employees, particularly in relation to safety, reliability, and trust.",
]

for text in body_paragraphs:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15

outcome = doc.add_paragraph()
outcome.paragraph_format.space_after = Pt(8)
run = outcome.add_run("Outcome: ")
run.bold = True
run.font.name = "Arial"
run.font.size = Pt(11)
run2 = outcome.add_run("[Insert sanction/outcome here]")
run2.font.name = "Arial"
run2.font.size = Pt(11)

appeal = doc.add_paragraph(
    "You are advised of your right to appeal this finding in accordance with company policy. Any appeal must be submitted in writing within [number] days of receipt of this letter and addressed to [appeal authority/title]."
)
appeal.paragraph_format.space_after = Pt(10)

closing = doc.add_paragraph("Yours faithfully,")
closing.paragraph_format.space_after = Pt(18)

sig = doc.add_paragraph()
sig.paragraph_format.space_after = Pt(2)
sig.add_run("[Name]").bold = True

sig2 = doc.add_paragraph("[Title]")
sig2.paragraph_format.space_after = Pt(0)

sig3 = doc.add_paragraph("[Company Name]")
sig3.paragraph_format.space_after = Pt(0)

doc.save(OUTPUT)
print(OUTPUT)
